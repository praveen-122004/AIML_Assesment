# backend/parsers/docx_parser.py

from docx import Document
import base64

def parse_docx(file):
    """Extract text, tables, and images from a DOCX file as HTML."""
    doc = Document(file)
    html = ""

    # Extract text paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            html += f"<p>{para.text}</p>"

    # Extract tables
    for table in doc.tables:
        html += "<table border='1'>"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{cell.text}</td>"
            html += "</tr>"
        html += "</table>"

    # Extract images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image_base64 = base64.b64encode(image_data).decode("utf-8")
            html += f"<img src='data:image/png;base64,{image_base64}' />"

    return html
